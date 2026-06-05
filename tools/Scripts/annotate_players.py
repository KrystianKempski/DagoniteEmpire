#!/usr/bin/env python3
"""
annotate_players.py — Kraina Możliwości
- Detects player characters from text colour in each .docx
- Adds "players": [...] to nodes in outputKraina/adventure_map.json
- Optionally inserts [CharacterName] before speaking blocks (black tag run)

Usage:
  python3 annotate_players.py           # backup + annotate docx + update JSON
  python3 annotate_players.py --dry-run
  python3 annotate_players.py --json-only   # only update adventure_map.json
"""

from __future__ import annotations

import json
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from kraina_colors import (
    KRAINA_BACKUP,
    KRAINA_JSON,
    KRAINA_ROOT,
    detect_heroes_in_paragraphs,
    para_dominant_color,
    resolve_hero,
    resolve_map_name,
)

TAG_RE = re.compile(r"^\s*\[[\w ]+\]")


def make_tag_run(tag: str):
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    r.insert(0, rpr)
    t = OxmlElement("w:t")
    t.text = f"[{tag}] "
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    return r


def insert_tag_before_first_run(para, tag: str) -> None:
    runs = para.runs
    if not runs:
        return
    first_r = runs[0]._r
    para._p.insert(list(para._p).index(first_r), make_tag_run(tag))


def annotate_document(docx_path: Path, dry_run: bool, tag_docx: bool) -> set[str]:
    doc = Document(str(docx_path))
    prev_char: str | None = None
    characters_found: set[str] = set()
    modified = False

    for para in doc.paragraphs:
        dom_color = para_dominant_color(para)

        if dom_color is None:
            prev_char = None
            continue

        char = resolve_hero(dom_color, para.text)

        if char is None:
            continue

        characters_found.add(char)

        if tag_docx and char != prev_char:
            if not dry_run and not TAG_RE.match(para.text):
                insert_tag_before_first_run(para, char)
                modified = True

        prev_char = char

    if modified:
        doc.save(str(docx_path))

    return characters_found


def update_json(results: dict[str, set[str]]) -> None:
    with open(KRAINA_JSON, encoding="utf-8") as f:
        data = json.load(f)

    node_map = {node["name"]: node for node in data.get("nodes", [])}
    node_names = set(node_map)

    updated, unmatched = 0, []
    for stem, chars in results.items():
        map_name = resolve_map_name(stem, node_names)
        if map_name:
            node_map[map_name]["players"] = sorted(chars)
            updated += 1
        else:
            unmatched.append(stem)

    with open(KRAINA_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\nJSON updated: {updated} document nodes")
    if unmatched:
        print(f"Unmatched local files (no JSON node): {len(unmatched)}")
        for s in unmatched[:15]:
            print(f"  - {s}")
        if len(unmatched) > 15:
            print(f"  … and {len(unmatched) - 15} more")


def _path_filter() -> str | None:
    for arg in sys.argv[1:]:
        if arg.startswith("--folder="):
            return arg.split("=", 1)[1]
    return None


def main() -> None:
    dry_run = "--dry-run" in sys.argv
    json_only = "--json-only" in sys.argv
    tag_docx = not json_only and not dry_run
    folder_filter = _path_filter()

    if not KRAINA_ROOT.is_dir():
        sys.exit(f"Brak folderu kampanii: {KRAINA_ROOT}")

    if tag_docx and not dry_run:
        if not KRAINA_BACKUP.exists():
            print(f"Creating backup at {KRAINA_BACKUP} …")
            shutil.copytree(str(KRAINA_ROOT), str(KRAINA_BACKUP))
            print("Backup done.\n")
        else:
            print(f"Backup already exists at {KRAINA_BACKUP} – skipping.\n")

    all_docx = sorted(KRAINA_ROOT.rglob("*.docx"))
    if folder_filter:
        all_docx = [p for p in all_docx if folder_filter in str(p.parent)]
    mode = "[DRY RUN] " if dry_run else ("[JSON ONLY] " if json_only else "")
    print(f"{mode}Processing {len(all_docx)} files …\n")

    results: dict[str, set[str]] = {}
    for docx_path in all_docx:
        stem = docx_path.stem
        if json_only or dry_run:
            doc = Document(str(docx_path))
            chars = detect_heroes_in_paragraphs(doc.paragraphs)
        else:
            chars = annotate_document(docx_path, dry_run=False, tag_docx=True)
        results[stem] = chars
        label = ", ".join(sorted(chars)) if chars else "(tylko MG / NPC)"
        print(f"  {stem}: {label}")

    if not dry_run:
        update_json(results)
    else:
        print("\n[DRY RUN] Pliki .docx i JSON bez zmian.")


if __name__ == "__main__":
    main()
