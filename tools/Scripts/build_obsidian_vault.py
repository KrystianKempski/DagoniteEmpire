#!/usr/bin/env python3
"""
build_obsidian_vault.py
Generuje gotowy do użycia Vault Obsidian z plików .docx kampanii RPG.

Uruchomienie:
    python3 build_obsidian_vault.py

Wynik: folder  tools/Resources/Wiki/  gotowy do otwarcia w Obsidian
       ("Open folder as vault")
"""

import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("Brak python-docx. Zainstaluj: pip install python-docx")

# ── Ścieżki ──────────────────────────────────────────────────────────────────
BASE = Path(__file__).resolve().parent.parent          # tools/
ADVENTURES_DIR = BASE / "Resources" / "W służbie Bonefire"
WORLD_DIR = BASE / "Resources" / "Opis świata i zasady"
JSON_PATH = BASE / "Scripts" / "output" / "adventure_map.json"
MERMAID_PATH = BASE / "Scripts" / "output" / "adventure_map.md"
VAULT_DIR = BASE / "Resources" / "Wiki"
PARTIES_CONFIG_PATH = BASE.parent.parent / "dagonite-wiki" / "content" / "_meta" / "wiki-parties.json"

# ── Konfiguracja dostępu (drużyny) ────────────────────────────────────────────
# Wyliczamy DOMYŚLNY blok `access:` dla generowanych stron. Reguły można potem
# nadpisać ręcznie we frontmatter lub w _meta/wiki-access-overrides.json.
PARTY_BY_CHARACTER: dict[str, list[str]] = {}
CHAR_CANON: dict[str, str] = {}
CAMPAIGN_PARTY_IDS: list[str] = []


def _norm_char(value: str) -> str:
    return re.sub(r"[\s\-_.]+", "", str(value)).lower()


def resolve_char(value: str) -> str | None:
    return CHAR_CANON.get(_norm_char(value))


def load_wiki_access_config() -> None:
    global CAMPAIGN_PARTY_IDS
    if not PARTIES_CONFIG_PATH.exists():
        print(f"  ↷ Brak {PARTIES_CONFIG_PATH} — access: zostanie ograniczony do prostych reguł")
        return
    try:
        cfg = json.loads(PARTIES_CONFIG_PATH.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError) as ex:
        print(f"  ! Nie udało się wczytać wiki-parties.json: {ex}")
        return

    for pid, party in cfg.get("parties", {}).items():
        for name in party.get("characters", []):
            PARTY_BY_CHARACTER.setdefault(name, []).append(pid)
            CHAR_CANON[_norm_char(name)] = name
    for canonical, aliases in cfg.get("characterAliases", {}).items():
        for alias in aliases:
            CHAR_CANON[_norm_char(alias)] = canonical

    ids: set[str] = set()
    for camp in cfg.get("campaigns", []):
        ids.update(camp.get("partyIds", []))
    CAMPAIGN_PARTY_IDS = sorted(ids)


def access_block_lines(visibility: str, characters: list[str] | None = None,
                       parties: list[str] | None = None) -> list[str]:
    lines = ["access:", f"  visibility: {visibility}"]
    if characters:
        lines.append("  characters: [" + ", ".join(yaml_str(c) for c in characters) + "]")
    if parties:
        lines.append("  parties: [" + ", ".join(parties) + "]")
    return lines


def scene_access(name: str, players: list[str]) -> tuple[str, list[str] | None, list[str] | None]:
    """Domyślny dostęp sceny: mirror heurystyki z build_wiki_access_manifest.py."""
    title = (name or "").strip()
    if title.lower().startswith("wstęp"):
        hero = re.sub(r"(?i)^wstęp\s*", "", title).strip()
        canonical = resolve_char(hero)
        if canonical:
            return ("characters", [canonical], None)
        return ("party", None, CAMPAIGN_PARTY_IDS)

    if not players:
        return ("party", None, CAMPAIGN_PARTY_IDS)

    canon = sorted({resolve_char(p) or p for p in players})
    if len(players) <= 2:
        return ("characters", canon, None)

    pids = sorted({pid for p in players for pid in PARTY_BY_CHARACTER.get(resolve_char(p) or p, [])})
    return ("party", None, pids or CAMPAIGN_PARTY_IDS)


# ── Pomocnicze ───────────────────────────────────────────────────────────────

def slugify(name: str) -> str:
    """Zamienia nazwę pliku na bezpieczny string (bez znaków spec.)."""
    name = name.strip()
    # usuń znaki, które Obsidian może mieć problem przetworzyć w linkach
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '', name)
    return name


def extract_text_from_docx(path: Path) -> str:
    """Wyciąga tekst z .docx jako Markdown (nagłówki, akapity, tabele)."""
    try:
        doc = Document(str(path))
    except Exception as e:
        return f"_[Błąd odczytu pliku: {e}]_"

    lines: list[str] = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        if "Heading 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style:
            lines.append(f"### {text}")
        elif "Title" in style:
            lines.append(f"# {text}")
        else:
            lines.append(text)

    # tabele
    for table in doc.tables:
        lines.append("")
        header_done = False
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if not header_done:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                header_done = True
        lines.append("")

    return "\n".join(lines)


def yaml_str(value: str) -> str:
    """Formatuje wartość jako string YAML (z cudzysłowami jeśli potrzeba)."""
    value = str(value).replace('"', '\\"')
    return f'"{value}"'


def make_frontmatter(node: dict[str, Any], players: list[str], gdrive: str) -> str:
    """Generuje blok YAML frontmatter dla notatki Obsidian."""
    title = yaml_str(node.get("name", ""))
    date_val = yaml_str(node.get("date", "brak info"))
    loc = yaml_str(node.get("location", "brak info"))
    folder = yaml_str(node.get("folder", ""))
    tags_list = ["przygoda"]
    folder_name = node.get("folder", "")
    if "Akt 1" in folder_name:
        tags_list.append("akt-1")
    elif "Akt 2" in folder_name:
        tags_list.append("akt-2")
    elif "Akt 3" in folder_name:
        tags_list.append("akt-3")
    elif "Akt 4" in folder_name:
        tags_list.append("akt-4")
    elif "WSPÓLNE" in folder_name:
        tags_list.append("wspolne")

    for p in players:
        tags_list.append(p.lower().replace(" ", "-"))

    tags_yaml = "[" + ", ".join(tags_list) + "]"
    players_yaml = "[" + ", ".join(yaml_str(p) for p in players) + "]"

    visibility, chars, parties = scene_access(node.get("name", ""), players)

    lines = [
        "---",
        f"title: {title}",
        f"ingame_date: {date_val}",
        f"location: {loc}",
        f"act: {folder}",
        f"players: {players_yaml}",
        f"gdrive: {yaml_str(gdrive)}",
        f"tags: {tags_yaml}",
        *access_block_lines(visibility, chars, parties),
        "---",
        "",
    ]
    return "\n".join(lines)


# ── Ładowanie danych z JSON ───────────────────────────────────────────────────

def load_map_data() -> tuple[dict, dict, list]:
    """Zwraca: nodes_by_name, nodes_by_id, edges."""
    with open(JSON_PATH, encoding="utf-8") as f:
        data = json.load(f)

    nodes_by_name: dict[str, dict] = {}
    nodes_by_id: dict[str, dict] = {}
    for node in data.get("nodes", []):
        name = node.get("name", "")
        nid = node.get("id", "")
        nodes_by_name[name] = node
        if nid:
            nodes_by_id[nid] = node

    edges = data.get("edges", [])
    return nodes_by_name, nodes_by_id, edges


def build_links_map(edges: list, nodes_by_id: dict) -> dict[str, list[str]]:
    """Buduje mapę: id → [nazwy scen, do których prowadzą linki]."""
    result: dict[str, list[str]] = {}
    for edge in edges:
        from_id = edge.get("from", "")
        to_id = edge.get("to", "")
        to_node = nodes_by_id.get(to_id)
        if to_node:
            result.setdefault(from_id, []).append(to_node.get("name", ""))
    return result


# ── Konwersja przygód ────────────────────────────────────────────────────────

# Folder, do którego trafiają wygenerowane sceny (scena = surowy transkrypt sesji).
# Wcześniej nazywał się "Przygody" → "Archiwum sesji" — teraz pod folderem kampanii.
CAMPAIGN_DIR_NAME = "W służbie Bonefyre"
SCENES_DIR_NAME = f"{CAMPAIGN_DIR_NAME}/Archiwum sesji"
CHARS_DIR_NAME = f"{CAMPAIGN_DIR_NAME}/Postacie"

def convert_adventures(nodes_by_name: dict, nodes_by_id: dict, links_map: dict) -> None:
    output_base = VAULT_DIR / SCENES_DIR_NAME

    for act_dir in sorted(ADVENTURES_DIR.iterdir()):
        if not act_dir.is_dir():
            continue
        act_name = act_dir.name
        act_out = output_base / act_name
        act_out.mkdir(parents=True, exist_ok=True)

        scene_files: list[str] = []

        for docx_file in sorted(act_dir.glob("*.docx")):
            # pomiń pliki tymczasowe LibreOffice
            if docx_file.name.startswith(".~"):
                continue

            stem = docx_file.stem
            md_name = slugify(stem) + ".md"
            out_path = act_out / md_name

            node = nodes_by_name.get(stem, {})
            node_id = node.get("id", "")
            players = node.get("players", [])
            gdrive = node.get("url", "")

            # frontmatter
            fm = make_frontmatter(node, players, gdrive)

            # treść
            body = extract_text_from_docx(docx_file)

            # wikilinki "Zobacz też"
            linked_names = links_map.get(node_id, [])
            see_also = ""
            if linked_names:
                see_also = "\n\n---\n\n## Zobacz też\n\n"
                see_also += "\n".join(f"- [[{slugify(n)}]]" for n in linked_names)

            content = fm + f"# {stem}\n\n" + body + see_also

            out_path.write_text(content, encoding="utf-8")
            scene_files.append(slugify(stem))

        # index aktu
        _write_act_index(act_out, act_name, scene_files)

    print(f"  ✓ {SCENES_DIR_NAME} → {output_base}")


def _write_act_index(act_out: Path, act_name: str, scene_files: list[str]) -> None:
    lines = [
        f"# {act_name}",
        "",
        "## Sceny",
        "",
    ]
    for sf in scene_files:
        lines.append(f"- [[{sf}]]")
    (act_out / "_Index.md").write_text("\n".join(lines), encoding="utf-8")


# ── Konwersja świata i zasad ──────────────────────────────────────────────────

def convert_world() -> None:
    world_out = VAULT_DIR / "Świat i zasady"
    world_out.mkdir(parents=True, exist_ok=True)

    # specjalne foldery
    for item in sorted(WORLD_DIR.iterdir()):
        if item.is_dir():
            sub_out = world_out / item.name
            sub_out.mkdir(exist_ok=True)
            for docx_file in sorted(item.glob("*.docx")):
                if docx_file.name.startswith(".~"):
                    continue
                _convert_world_file(docx_file, sub_out)
        elif item.suffix.lower() == ".docx":
            _convert_world_file(item, world_out)

    print(f"  ✓ Świat i zasady → {world_out}")


def _convert_world_file(docx_file: Path, out_dir: Path) -> None:
    stem = docx_file.stem
    md_name = slugify(stem) + ".md"
    out_path = out_dir / md_name

    # kategoryzacja po tytule
    if any(k in stem.lower() for k in ["instrukc", "zasad", "przykład"]):
        tags = "[zasady]"
    elif any(k in stem.lower() for k in ["panteon", "zakon", "religia"]):
        tags = "[lore, religia]"
    elif any(k in stem.lower() for k in ["geopolit", "ludy", "rody"]):
        tags = "[lore, geopolityka]"
    elif any(k in stem.lower() for k in ["cennik", "menu", "ceny"]):
        tags = "[lore, ekonomia]"
    else:
        tags = "[lore]"

    # "Świat i zasady" to publiczne lore (anonymousPublicPrefixes) → dostęp public.
    access = "\n".join(access_block_lines("public"))
    fm = f"---\ntitle: \"{stem}\"\ntags: {tags}\n{access}\n---\n\n"
    body = extract_text_from_docx(docx_file)
    content = fm + f"# {stem}\n\n" + body

    out_path.write_text(content, encoding="utf-8")


# ── Strony postaci ────────────────────────────────────────────────────────────

# Lista postaci graczy z annotate_players.py
CHARACTERS = [
    "Werner", "Lawenda", "Sariel", "Dorian", "Tomin",
    "Udar", "Sharu", "Umbra", "Sir Bron", "Sir Cedrick", "Baron Mevir",
]

def create_character_pages(nodes_by_name: dict, nodes_by_id: dict, links_map: dict) -> None:
    chars_out = VAULT_DIR / CHARS_DIR_NAME
    chars_out.mkdir(parents=True, exist_ok=True)

    # zbierz sceny, w których postać bierze udział
    char_scenes: dict[str, list[str]] = {c: [] for c in CHARACTERS}
    for node in nodes_by_name.values():
        players = node.get("players", [])
        for p in players:
            if p in char_scenes:
                char_scenes[p].append(node.get("name", ""))

    for char, scenes in char_scenes.items():
        canonical = resolve_char(char) or char
        lines = [
            "---",
            f"title: {yaml_str(char)}",
            *access_block_lines("characters", [canonical]),
            "---",
            "",
            f"# {char}",
            "",
            "> Strona postaci gracza.",
            "",
            "## Sceny",
            "",
        ]
        for s in sorted(set(scenes)):
            lines.append(f"- [[{slugify(s)}]]")

        (chars_out / f"{char}.md").write_text("\n".join(lines), encoding="utf-8")

    print(f"  ✓ Postacie → {chars_out}")


# ── Strona główna ──────────────────────────────────────────────────────────────

def create_home_page(nodes_by_name: dict) -> None:
    # Policz sceny per akt
    act_counts: dict[str, int] = {}
    for node in nodes_by_name.values():
        folder = node.get("folder", "")
        if folder and folder != "<root>":
            act_counts[folder] = act_counts.get(folder, 0) + 1

    lines = [
        "# 🎲 DagoniteEmpire — Wiki Kampanii",
        "",
        "> Witaj w encyklopedii kampanii **W służbie Bonefire**.",
        "> Kliknij na folder w panelu bocznym lub użyj poniższych linków.",
        "",
        "## Nawigacja",
        "",
        "| Sekcja | Opis |",
        "| --- | --- |",
        f"| [[{SCENES_DIR_NAME}/_Index]] | Surowy transkrypt scen (archiwum) |",
        "| [[Świat i zasady/_Index]] | Opis świata, zasady, geopolityka |",
        "| [[Postacie/_Index]] | Karty postaci graczy |",
        "| [[Mapy/Mapa powiązań]] | Graf powiązań między scenami |",
        "",
        "## Akty kampanii",
        "",
    ]

    # sortuj akty
    act_order = ["Akt 1", "Akt 2", "Akt 3", "Akt 4", "WSPÓLNE"]
    sorted_acts = sorted(
        act_counts.items(),
        key=lambda x: next((i for i, k in enumerate(act_order) if k in x[0]), 99)
    )
    for act, count in sorted_acts:
        act_slug = slugify(act)
        lines.append(f"- **[[{SCENES_DIR_NAME}/{act_slug}/_Index|{act}]]** — {count} scen")

    lines += [
        "",
        "## Mapa powiązań",
        "",
        "![[Mapy/Mapa powiązań]]",
        "",
    ]

    home_path = VAULT_DIR / "Home.md"
    if home_path.exists():
        # Ręcznie edytowane Home.md nie jest nadpisywane.
        # Usuń plik (lub ustaw FORCE_HOME=1), aby wymusić regenerację.
        import os
        if not os.environ.get("FORCE_HOME"):
            print(f"  ↷ Home.md istnieje — pomijam (ustaw FORCE_HOME=1, aby nadpisać)")
            return
    home_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  ✓ Home.md → {home_path}")


# ── Mapa Mermaid ──────────────────────────────────────────────────────────────

def copy_mermaid_map() -> None:
    maps_out = VAULT_DIR / "Mapy"
    maps_out.mkdir(exist_ok=True)
    dest = maps_out / "Mapa powiązań.md"

    raw = MERMAID_PATH.read_text(encoding="utf-8")
    header = "---\ntitle: \"Mapa powiązań między scenami\"\ntags: [mapa]\n---\n\n"
    dest.write_text(header + raw, encoding="utf-8")
    print(f"  ✓ Mapa Mermaid → {dest}")


# ── Index stron ───────────────────────────────────────────────────────────────

def create_index_pages() -> None:
    # Archiwum sesji/_Index
    adv_dir = VAULT_DIR / SCENES_DIR_NAME
    adv_acts = sorted(d.name for d in adv_dir.iterdir() if d.is_dir()) if adv_dir.exists() else []
    adv_lines = [
        "---",
        'title: "4. Archiwum sesji"',
        "tags: [index, archiwum]",
        "---",
        "",
        "# Archiwum sesji — przegląd\n",
    ]
    for act in adv_acts:
        adv_lines.append(f"## [[{act}/_Index|{act}]]\n")
    (adv_dir / "_Index.md").write_text("\n".join(adv_lines), encoding="utf-8")

    # Świat i zasady/_Index
    world_dir = VAULT_DIR / "Świat i zasady"
    world_dir.mkdir(exist_ok=True)
    world_files = sorted(world_dir.rglob("*.md"))
    world_lines = [
        "---",
        'title: "1. Świat i zasady"',
        "tags: [index, swiat]",
        "---",
        "",
        "# Świat i zasady — przegląd\n",
    ]
    for wf in world_files:
        rel = wf.relative_to(world_dir)
        stem = wf.stem
        if stem == "_Index":
            continue
        world_lines.append(f"- [[{rel.with_suffix('')}|{stem}]]")
    (world_dir / "_Index.md").write_text("\n".join(world_lines), encoding="utf-8")

    # Postacie/_Index (pod kampanią)
    chars_dir = VAULT_DIR / CHARS_DIR_NAME
    chars_dir.mkdir(parents=True, exist_ok=True)
    chars_lines = [
        "---",
        'title: "1. Postacie"',
        "tags: [index, postacie]",
        "---",
        "",
        "# Postacie graczy\n",
    ]
    for char in CHARACTERS:
        chars_lines.append(f"- [[{char}]]")
    (chars_dir / "_Index.md").write_text("\n".join(chars_lines), encoding="utf-8")

    print(f"  ✓ Strony indeksowe")


# ── Konfiguracja Obsidian ────────────────────────────────────────────────────

def write_obsidian_config() -> None:
    obsidian_dir = VAULT_DIR / ".obsidian"
    obsidian_dir.mkdir(exist_ok=True)

    # app.json — ustawia Home.md jako stronę startową, włącza Mermaid
    app_config = {
        "defaultViewMode": "preview",
        "newFileLocation": "current",
        "attachmentFolderPath": "Zasoby",
        "legacyEditor": False,
    }
    (obsidian_dir / "app.json").write_text(
        json.dumps(app_config, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # graph.json — konfiguracja widoku grafu
    graph_config = {
        "collapse-filter": False,
        "search": "",
        "showTags": True,
        "showAttachments": False,
        "hideUnresolved": False,
        "showOrphans": True,
        "collapse-color-groups": False,
        "colorGroups": [
            {"query": "tag:akt-1", "color": {"a": 1, "rgb": 14737632}},
            {"query": "tag:akt-2", "color": {"a": 1, "rgb": 5614830}},
            {"query": "tag:akt-3", "color": {"a": 1, "rgb": 16740352}},
            {"query": "tag:akt-4", "color": {"a": 1, "rgb": 16711680}},
        ],
    }
    (obsidian_dir / "graph.json").write_text(
        json.dumps(graph_config, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # core-plugins.json
    plugins = ["graph", "backlink", "tag-pane", "search",
               "template", "note-composer", "command-palette"]
    (obsidian_dir / "core-plugins.json").write_text(
        json.dumps(plugins, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(f"  ✓ Konfiguracja Obsidian → {obsidian_dir}")


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    print(f"\n🏗️  Budowanie Obsidian Vault w: {VAULT_DIR}\n")
    VAULT_DIR.mkdir(parents=True, exist_ok=True)

    print("🔐 Wczytywanie konfiguracji dostępu (drużyny)...")
    load_wiki_access_config()

    print("📖 Wczytywanie mapy przygód...")
    nodes_by_name, nodes_by_id, edges = load_map_data()
    links_map = build_links_map(edges, nodes_by_id)
    print(f"   {len(nodes_by_name)} węzłów, {len(edges)} krawędzi")

    print("\n📜 Konwertowanie przygód...")
    convert_adventures(nodes_by_name, nodes_by_id, links_map)

    print("\n🌍 Konwertowanie opisu świata i zasad...")
    convert_world()

    print("\n🧙 Tworzenie stron postaci...")
    create_character_pages(nodes_by_name, nodes_by_id, links_map)

    print("\n🗺️  Kopiowanie mapy Mermaid...")
    copy_mermaid_map()

    print("\n📋 Tworzenie stron indeksowych...")
    create_index_pages()

    print("\n🏠 Tworzenie strony głównej...")
    create_home_page(nodes_by_name)

    print("\n⚙️  Zapisywanie konfiguracji Obsidian...")
    write_obsidian_config()

    # Podsumowanie
    md_count = len(list(VAULT_DIR.rglob("*.md")))
    print(f"\n✅ Gotowe! Wygenerowano {md_count} plików .md")
    print(f"   Vault: {VAULT_DIR}")
    print("\n📌 Następne kroki:")
    print("   1. Zainstaluj Obsidian desktop: https://obsidian.md/download")
    print("   2. Otwórz Obsidian → 'Open folder as vault'")
    print(f"   3. Wskaż folder: {VAULT_DIR}")
    print("   4. Opcjonalnie: zainstaluj wtyczkę 'Dataview' do zaawansowanego filtrowania")
    print("   5. Opcjonalnie: Quartz (https://quartz.jzhao.xyz/) → bezpłatna strona web\n")


if __name__ == "__main__":
    main()
