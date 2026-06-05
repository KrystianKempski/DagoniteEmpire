#!/usr/bin/env python3
"""
Generate Archiwum sesji/*.md from local Kraina .docx into dagonite-wiki/content.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

from build_kraina_wiki_tags import act_number_from_path, load_map_players, players_for_docx, tags_for_scene
from kraina_colors import KRAINA_JSON, KRAINA_ROOT, resolve_map_name
from kraina_header import extract_header_from_docx, parse_header

SCRIPT_DIR = Path(__file__).resolve().parent
WIKI_ROOT = SCRIPT_DIR.parent.parent.parent / "dagonite-wiki" / "content" / "Kraina Możliwości"
ARCHIVE = WIKI_ROOT / "Archiwum sesji"

MECHANICS_RE = re.compile(
    r"^\s*(\(test |\[test |\d+d\d|vs |HP |AC |rzut |throw |save )",
    re.IGNORECASE,
)
BRACKET_RE = re.compile(r"^\s*\[[\w ]+\]")


def slug_filename(stem: str) -> str:
    s = stem.replace('"', "").replace("'", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s + ".md"


def act_folder_name(act_dir: Path) -> str:
    return act_dir.name  # e.g. "Akt 1 Nowi bohaterowie w Warrington"


def extract_narrative(docx: Path, max_chars: int = 12000) -> str:
    doc = Document(docx)
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t or MECHANICS_RE.match(t) or BRACKET_RE.match(t):
            continue
        if t.startswith("(test"):
            continue
        parts.append(t)
    text = "\n\n".join(parts)
    if len(text) > max_chars:
        text = text[: max_chars - 3].rsplit("\n\n", 1)[0] + "…"
    return text


def yaml_tags(tags: list[str]) -> str:
    return "[" + ", ".join(tags) + "]"


def scene_summary_block(text: str, max_paras: int = 8) -> str:
    """First few narrative paragraphs as session summary."""
    paras = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 40]
    chosen = paras[:max_paras]
    if not chosen:
        return "_Brak streszczenia — uzupełnij po sesji._"
    return "\n\n".join(chosen)


def write_scene(docx: Path, act_dir: Path, dry_run: bool = False) -> None:
    act = act_number_from_path(act_dir)
    map_players = load_map_players()
    players = players_for_docx(docx, map_players)
    tags = tags_for_scene(players, act)

    doc = Document(docx)
    header_line = extract_header_from_docx(doc.paragraphs[:8])
    date, location = parse_header(docx.stem, header_line)
    if date == "brak info":
        date = "nieznana"
    if location == "brak info":
        location = "nieznana"

    narrative = extract_narrative(docx)
    title = docx.stem

    out_dir = ARCHIVE / act_folder_name(act_dir)
    out_path = out_dir / slug_filename(docx.stem)

    players_line = ", ".join(players) if players else "—"
    body = f"""---
title: "{title}"
tags: {yaml_tags(tags)}
date: "{date}"
location: "{location}"
players: [{", ".join(f'"{p}"' for p in players)}]
---

# {title}

> **{date}** · {location}  
> **Obecni:** {players_line}

## Streszczenie

{scene_summary_block(narrative)}

## Notatki MG

_Surowy zapis sesji w archiwum kampanii (Google Drive / .docx)._
"""

    if dry_run:
        print(out_path)
        return
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path.write_text(body, encoding="utf-8")


def main() -> None:
    dry = "--dry-run" in sys.argv
    for act_dir in sorted(KRAINA_ROOT.glob("Akt *")):
        for docx in sorted(act_dir.glob("*.docx")):
            write_scene(docx, act_dir, dry_run=dry)
    print(f"Done → {ARCHIVE}")


if __name__ == "__main__":
    main()
